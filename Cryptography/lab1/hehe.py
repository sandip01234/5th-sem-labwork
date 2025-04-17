from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('HTML Assignment', 0)

# Add questions and answers to the document
questions_answers = [
    ("1. Difference between HTML and XHTML", """
        | Feature                | HTML                                      | XHTML                                    |
        |------------------------|-------------------------------------------|------------------------------------------|
        | Syntax                 | Less strict (optional closing tags)       | Strict syntax (requires closing tags)   |
        | Tag case               | Tags are case-insensitive (can be lowercase/uppercase) | Tags must be written in lowercase     |
        | Doctype Declaration    | Less specific (e.g., <!DOCTYPE HTML>)      | More specific (e.g., <!DOCTYPE html PUBLIC...>) |
        | Attribute Quotes       | Attribute values do not always need quotes | Attribute values must always be in quotes |
        | Error Handling         | Browsers try to render even with errors    | Errors must be fixed to display properly |
        | Document Structure     | Tolerant of poorly structured code        | Requires well-structured code (well-formed) |
        | Compatibility          | Compatible with all browsers              | Requires modern browsers for full support |
    """),
    ("2. Explain the history of HTML in brief", """
        HTML (Hypertext Markup Language) was created by Tim Berners-Lee in 1991 as a system for formatting and displaying documents on the World Wide Web. The first version of HTML was basic, allowing the creation of simple web pages with text and links. Over time, HTML evolved through various versions:
        - HTML 2.0 (1995): Standardized the basic elements and attributes of HTML.
        - HTML 3.2 (1997): Introduced tables, styles, and applets.
        - HTML 4.01 (1999): Improved accessibility, scripting, and multimedia support.
        - HTML5 (2014): A significant update that introduced new elements, APIs, and improvements for mobile web development and multimedia.
    """),
    ("3. What is HTML attribute? Describe HTML elements with its types", """
        - HTML Attribute: An attribute provides additional information about an HTML element. It is placed within the opening tag and usually comes in name-value pairs.
        Example: <a href="https://example.com">Link</a>

        - HTML Elements:
          1. Block-level elements: These elements take up the full width available and start on a new line (e.g., <div>, <h1>).
          2. Inline elements: These elements take up only as much width as necessary and do not start on a new line (e.g., <span>, <a>).
          3. Empty (self-closing) elements: Elements that do not have closing tags (e.g., <img>, <br>).
    """),
    ("4. What is tag in HTML? Describe the different types of Tags in HTML", """
        - Tag: A tag is a special keyword enclosed in angle brackets (< >) that defines an element in an HTML document. Tags are used to structure content on the web page.
        Example: <h1>Title</h1>

        - Types of Tags:
          1. Opening Tags: Defines the beginning of an element (e.g., <h1>).
          2. Closing Tags: Defines the end of an element (e.g., </h1>).
          3. Self-closing Tags: These tags do not have closing counterparts (e.g., <img />).
    """),
    ("5. Describe the use of hyperlink tag", """
        The hyperlink tag <a> is used to create links between web pages or to external websites. The href attribute specifies the destination URL.
        Example:
        <a href="https://www.example.com">Visit Example</a>
    """),
    ("6. Create a basic HTML table with cell padding, cellspacing, rowspan, colspan and border attribute", """
        <table border="1" cellspacing="10" cellpadding="15">
          <tr>
            <td rowspan="2">Rowspan Example</td>
            <td colspan="2">Colspan Example</td>
          </tr>
          <tr>
            <td>Cell 1</td>
            <td>Cell 2</td>
          </tr>
        </table>
    """),
    ("7. Create a basic HTML form with various input fields", """
        <form action="/submit" method="POST">
          <label for="name">Name:</label>
          <input type="text" id="name" name="name"><br>

          <label for="age">Age:</label>
          <input type="number" id="age" name="age"><br>

          <label for="salary">Salary:</label>
          <input type="number" id="salary" name="salary"><br>

          <label for="select">Select a value:</label>
          <select id="select" name="select">
            <option value="option1">Option 1</option>
            <option value="option2">Option 2</option>
          </select><br>

          <label for="gender">Gender:</label>
          <input type="radio" id="male" name="gender" value="male">
          <label for="male">Male</label>
          <input type="radio" id="female" name="gender" value="female">
          <label for="female">Female</label><br>

          <label for="hobbies">Hobbies:</label>
          <input type="checkbox" id="reading" name="hobbies" value="reading">
          <label for="reading">Reading</label>
          <input type="checkbox" id="travelling" name="hobbies" value="travelling">
          <label for="travelling">Travelling</label><br>

          <label for="description">Description:</label><br>
          <textarea id="description" name="description"></textarea><br>

          <input type="submit" value="Submit">
        </form>
    """),
    ("8. Discuss HTML events with examples", """
        HTML events are actions that occur in response to user interactions. These events can trigger JavaScript functions or other actions.
        - Common HTML events:
          1. onclick: Triggered when an element is clicked.
             Example: <button onclick="alert('Hello!')">Click Me</button>
          2. onmouseover: Triggered when the mouse pointer hovers over an element.
             Example: <div onmouseover="this.style.backgroundColor='yellow'">Hover over me</div>
          3. onchange: Triggered when the value of an element changes.
             Example: <input type="text" onchange="alert('Value changed')">
          4. onfocus: Triggered when an element gains focus.
             Example: <input type="text" onfocus="this.style.backgroundColor='lightblue'">
    """),
]

# Add content to document
for question, answer in questions_answers:
    doc.add_heading(question, level=2)
    doc.add_paragraph(answer)

# Save the document
file_path = "C:/Users/HP VICTUS/OneDrive/Desktop/HTML_Assignment.docx"

doc.save(file_path)

file_path
